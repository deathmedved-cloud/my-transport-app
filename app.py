import io
import json
import os
import tempfile
import subprocess
import mimetypes
import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

st.set_page_config(page_title="Автоматизация Документов", layout="wide")
st.title("🚚 Формирование Актов и Счетов по шаблону")

# --- НАСТРОЙКА КЛЮЧА GEMINI ---
with st.sidebar:
    st.header("⚙️ Настройки AI")
    default_key = st.secrets.get("GEMINI_API_KEY", "")
    api_key = st.text_input("Gemini API Key", value=default_key, type="password")

# --- ВСПОМОГАТЕЛЬНАЯ ФУНКЦИЯ ДЛЯ НАСТРОЙКИ ШРИФТА ---
def apply_global_font(doc, font_name="Times New Roman", size_pt=12):
    style = doc.styles['Normal']
    font = style.font
    font.name = font_name
    font.size = Pt(size_pt)

# --- РАСПОЗНАВАНИЕ СКАНО / ДОКУМЕНТОВ ---
st.subheader("1. Загрузка исходной заявки / документа")
uploaded_file = st.file_uploader("Загрузите файл (PDF, фото, DOCX и др.)")

if "doc_data" not in st.session_state:
    st.session_state.doc_data = {
        "doc_num": "25",
        "app_num": "б/н",
        "app_date": "06.08.2026",
        "city": "г. Смоленск",
        "doc_date": "10.08.2026",
        "customer_name": "ИП Саук Д.М.",
        "customer_director": "директора Саук Д.М.",
        "customer_basis": "свидетельства о государственной регистрации 590625501",
        "transport_details": "ТС Р712НХ67 / АМ385467",
        "route": "РБ,ТЛЦ №2 г. Брест - РФ, Московская обл.г. Реутов",
        "amount_num": "1659",
        "amount_words": "одна тысяча шестьсот пятьдесят девять",
        "currency": "евро",
        "nds": "0%",
        "payment_note": "Оплата в рублях РФ по курсу НБ РБ на дату оплаты.",
        "customer_address_unp": "230026, РБ г. Гродно\nул. Пестрака, 22-1\nУНП 590625501",
        "customer_signer": "Саук Д.М."
    }

if st.button("🤖 Извлечь данные для Акта"):
    final_key = api_key or default_key
    if not final_key:
        st.error("Укажите Gemini API Key!")
    elif not uploaded_file:
        st.warning("Сначала выберите файл для загрузки.")
    else:
        try:
            genai.configure(api_key=final_key)
            model = genai.GenerativeModel(
                model_name="gemini-3.6-flash",
                generation_config={"temperature": 0.0}
            )

            with st.spinner("Считываем переменные поля из документа..."):
                file_bytes = uploaded_file.getvalue()
                mime_type = uploaded_file.type or mimetypes.guess_type(uploaded_file.name)[0] or "application/octet-stream"

                prompt = """
                Внимательно изучи документ и извлеки ИСКЛЮЧИТЕЛЬНО следующие изменяемые переменные для Акта выполненных работ.
                Верни СТРОГО JSON-объект без markdown-разметки:

                {
                    "doc_num": "номер акта (напр. 25)",
                    "app_num": "номер договора/заявки (напр. б/н или 12)",
                    "app_date": "дата договора/заявки (напр. 06.08.2026)",
                    "city": "город составления (напр. г. Смоленск)",
                    "doc_date": "дата акта (напр. 10.08.2026)",
                    "customer_name": "название Заказчика (напр. ИП Саук Д.М.)",
                    "customer_director": "в лице кого (напр. директора Саук Д.М.)",
                    "customer_basis": "на основании чего (напр. свидетельства о государственной регистрации 590625501)",
                    "transport_details": "ТС и прицеп (напр. ТС Р712НХ67 / АМ385467)",
                    "route": "маршрут (напр. РБ,ТЛЦ №2 г. Брест - РФ, Московская обл.г. Реутов)",
                    "amount_num": "сумма цифрами (напр. 1659)",
                    "amount_words": "сумма прописью (напр. одна тысяча шестьсот пятьдесят девять)",
                    "currency": "валюта (напр. евро)",
                    "nds": "ставка НДС (напр. 0%)",
                    "payment_note": "примечание по оплате/курсу (напр. Оплата в рублях РФ по курсу НБ РБ на дату оплаты.)",
                    "customer_address_unp": "адрес и УНП/ИНН Заказчика для реквизитов",
                    "customer_signer": "ФИО подписывающего Заказчика (напр. Саук Д.М.)"
                }
                """

                response = model.generate_content([{"mime_type": mime_type, "data": file_bytes}, prompt])
                raw_text = response.text.strip().replace("```json", "").replace("```", "")
                parsed_json = json.loads(raw_text)
                st.session_state.doc_data.update(parsed_json)
                st.success("Данные успешно извлечены!")

        except Exception as e:
            st.error(f"Ошибка распознавания: {e}")

# --- ФОРМА РЕДАКТИРОВАНИЯ ТОЛЬКО ПЕРЕМЕННЫХ ПОЛЕЙ ---
st.subheader("2. Переменные поля Акта (изменяемая часть)")

data = st.session_state.doc_data
col1, col2 = st.columns(2)

with col1:
    doc_num = st.text_input("№ Акта и Счета", value=data.get("doc_num", "25"))
    app_num = st.text_input("№ договора/заявки", value=data.get("app_num", "б/н"))
    app_date = st.text_input("Дата договора/заявки", value=data.get("app_date", "06.08.2026г."))
    city = st.text_input("Город", value=data.get("city", "г. Смоленск"))
    doc_date = st.text_input("Дата Акта и Счета", value=data.get("doc_date", "10.08.2026г."))
    
    st.markdown("---")
    customer_name = st.text_input("Заказчик", value=data.get("customer_name", "ИП Саук Д.М."))
    customer_director = st.text_input("В лице", value=data.get("customer_director", "директора Саук Д.М."))
    customer_basis = st.text_input("Действующего на основании", value=data.get("customer_basis", "свидетельства о государственной регистрации 590625501"))

with col2:
    transport_details = st.text_input("ТС / авто", value=data.get("transport_details", "ТС Р712НХ67 / АМ385467"))
    route = st.text_input("Маршрут", value=data.get("route", "РБ,ТЛЦ №2 г. Брест  -  РФ, Московская обл.г. Реутов"))
    amount_num = st.text_input("Сумма цифрами", value=data.get("amount_num", "1659"))
    amount_words = st.text_input("Сумма прописью", value=data.get("amount_words", "одна тысяча шестьсот пятьдесят девять"))
    currency = st.text_input("Валюта", value=data.get("currency", "евро"))
    nds = st.text_input("НДС", value=data.get("nds", "0%"))
    payment_note = st.text_input("Условие оплаты", value=data.get("payment_note", "Оплата в рублях РФ по курсу НБ РБ на дату оплаты."))
    
    st.markdown("---")
    customer_address_unp = st.text_area("Адрес и УНП Заказчика", value=data.get("customer_address_unp", "230026, РБ г. Гродно\nул. Пестрака, 22-1\nУНП 590625501"))
    customer_signer = st.text_input("Подпись Заказчика", value=data.get("customer_signer", "Саук Д.М."))

# --- СТРОГАЯ ГЕНЕРАЦИЯ АКТА ПО ШАБЛОНУ ---
def build_act_doc():
    doc = Document()
    apply_global_font(doc, "Times New Roman", 12)
    
    # Шапка акта
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = h.add_run(f"Акт выполненных работ № {doc_num}\n об оказании транспортных услуг \nпо договору {app_num} от {app_date}")
    r1.bold = True
    r1.font.size = Pt(12)
    
    # Город и дата
    p_meta = doc.add_paragraph()
    p_meta.add_run(f"{city}\t\t\t\t\t\t\t\t\t\t\t{doc_date}")
    
    # Преамбула (статический текст + подстановка переменных)
    p_preamble = doc.add_paragraph()
    p_preamble.paragraph_format.first_line_indent = Inches(0.3)
    p_preamble.add_run(
        f"  {customer_name}, в лице {customer_director} действующего на основании {customer_basis}, "
        f"именуемое в дальнейшем Заказчик с одной стороны, и ООО«АВРОРА-ТРАНЗИТ», в лице директора Галенда С.В., "
        f"действующего на основании Устава, именуемое в дальнейшем Исполнитель с другой стороны, составили настоящий акт о нижеследующем:"
    )
    
    # Абзац 1
    p_b1 = doc.add_paragraph()
    p_b1.paragraph_format.first_line_indent = Inches(0.3)
    p_b1.add_run(
        f"  В установленные Договором сроки ООО«АВРОРА-ТРАНЗИТ» оказало         транспортные услуги по выполнению международной перевозки груза, "
        f"{transport_details},по маршруту: {route}."
    )
    
    # Абзац 2
    p_b2 = doc.add_paragraph()
    p_b2.paragraph_format.first_line_indent = Inches(0.3)
    p_b2.add_run(
        "  На основании изложенного Заказчик и Исполнитель заявляют, что оказанные транспортные услуги выполнены в полном объеме,и в срок. "
        "Заказчик претензий по объёму, качеству и срокам оказания услуг не имеет."
    )
    
    # Абзац 3 (Реквизиты оплаты)
    p_pay = doc.add_paragraph()
    p_pay.add_run(
        f"Указанную в договоре-заявке сумму  в размере {amount_num} ({amount_words}) {currency}, в том числе НДС {nds} следует перечислить ООО«АВРОРА-ТРАНЗИТ» по следующим реквизитам:\n"
        "Номер счёта: 40702810901130005079\n"
        "Валюта: RUR\n"
        "Банк: АО \"АЛЬФА-БАНК\"\n"
        "ИНН банка: 7728168971\n"
        "БИК: 044525593\n"
        "Адрес банка: 214004, Смоленская обл., г. Смоленск, ул. Николаева, д.8.\n"
        f"{payment_note}"
    )
    
    doc.add_paragraph("\n")
    
    # Подписи
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = f"{customer_name}\n{customer_address_unp}"
    table.rows[0].cells[1].text = "ООО «АВРОРА-ТРАНЗИТ»\n214022,РФ г. Смоленск\nул. Карбышева д. 15А,\nстр.2, помещ. К 52\nОГРН 1266700001501"
    
    table.rows[1].cells[0].text = f"\n\n                                  {customer_signer}"
    table.rows[1].cells[1].text = "\n\n                                                               Галенда С.В."
    
    # Явное приведение шрифта таблиц к Times New Roman 12pt
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.style = doc.styles['Normal']

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# --- ГЕНЕРАЦИЯ СЧЕТА НА ОСНОВЕ ДАННЫХ АКТА ---
def build_invoice_doc():
    doc = Document()
    apply_global_font(doc, "Times New Roman", 12)
    
    # Шапка банка (Таблица 2x2)
    table_bank = doc.add_table(rows=2, cols=2)
    table_bank.style = 'Table Grid'
    
    table_bank.rows[0].cells[0].text = "АО \"АЛЬФА-БАНК\"\n\nБанк получателя"
    table_bank.rows[0].cells[1].text = "БИК  044525593\nСч. №  30101810200000000593"
    table_bank.rows[1].cells[0].text = "ИНН 6700042504   КПП 670001001\n\nООО \"АВРОРА-ТРАНЗИТ\"\n\nПолучатель"
    table_bank.rows[1].cells[1].text = "Сч. №  40702810901130005079"

    for row in table_bank.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.style = doc.styles['Normal']

    doc.add_paragraph()
    
    h = doc.add_paragraph(f"Счет на оплату № {doc_num} от {doc_date}")
    h.runs[0].bold = True
    h.runs[0].font.size = Pt(12)

    p_party = doc.add_paragraph()
    p_party.add_run("Исполнитель: ").bold = True
    p_party.add_run("ООО \"АВРОРА-ТРАНЗИТ\", 214022, Смоленская область, г. о. город Смоленск, г. Смоленск, ул. Карбышева, д. 15А, стр. 2, помещ. К 52\n")
    p_party.add_run("Заказчик: ").bold = True
    p_party.add_run(f"{customer_name}, {customer_address_unp.replace(chr(10), ', ')}\n")
    p_party.add_run("Комментарий: ").bold = True
    p_party.add_run(f"Договор-заявка {app_num} от {app_date}\n{payment_note}")

    t_services = doc.add_table(rows=2, cols=7)
    t_services.style = 'Table Grid'
    
    headers = ["№", "Название услуги", "Кол-во", "Ед.изм.", "Цена", "НДС", "Сумма"]
    for i, title in enumerate(headers):
        t_services.rows[0].cells[i].text = title
        
    row = t_services.rows[1].cells
    row[0].text = "1"
    row[1].text = route
    row[2].text = "1"
    row[3].text = "ШТ."
    row[4].text = amount_num
    row[5].text = nds
    row[6].text = amount_num

    for r in t_services.rows:
        for cell in r.cells:
            for p in cell.paragraphs:
                p.style = doc.styles['Normal']

    doc.add_paragraph()
    
    p_total_text = doc.add_paragraph()
    p_total_text.add_run(f"Всего наименований 1 на сумму {amount_num} {currency}\n").bold = True
    p_total_text.add_run(f"{amount_words} {currency}").italic = True

    p_summary = doc.add_paragraph()
    p_summary.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_summary.add_run(f"Итого:  {amount_num}\nСумма НДС:  {nds}\nВсего к оплате:  {amount_num}").bold = True

    doc.add_paragraph("\n\n(должность) _________________ (подпись) _________________ (расшифровка подписи)")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# --- КОНВЕРТАТОР В PDF / TXT / DOCX ---
def convert_doc(doc_buf, target_fmt):
    docx_bytes = doc_buf.getvalue()
    if target_fmt == "PDF":
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_in:
                tmp_in.write(docx_bytes)
                tmp_in_path = tmp_in.name
            out_dir = tempfile.gettempdir()
            subprocess.run(["soffice", "--headless", "--convert-to", "pdf", tmp_in_path, "--outdir", out_dir], check=True)
            pdf_path = os.path.join(out_dir, os.path.splitext(os.path.basename(tmp_in_path))[0] + ".pdf")
            with open(pdf_path, "rb") as f:
                res_bytes = f.read()
            return res_bytes, "application/pdf", "pdf"
        except Exception:
            st.error("Для экспорта PDF добавьте 'libreoffice' в файл packages.txt.")
            return docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx"
            
    elif target_fmt == "TXT":
        doc = Document(io.BytesIO(docx_bytes))
        lines = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(lines).encode('utf-8'), "text/plain", "txt"
    else:
        return docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx"

# --- СКАЧИВАНИЕ ГОТОВЫХ ДОКУМЕНТОВ ---
st.subheader("3. Скачать сформированный Акт и Счет")

output_format = st.selectbox("Выберите формат выгрузки:", ["DOCX (.docx)", "PDF (.pdf)", "TXT (.txt)"])
fmt_code = output_format.split()[0]

c_btn1, c_btn2 = st.columns(2)

act_data, act_mime, act_ext = convert_doc(build_act_doc(), fmt_code)
invoice_data, inv_mime, inv_ext = convert_doc(build_invoice_doc(), fmt_code)

with c_btn1:
    st.download_button(
        label=f"📝 Скачать Акт ({fmt_code})",
        data=act_data,
        file_name=f"Акт_№_{doc_num}_от_{doc_date}.{act_ext}",
        mime=act_mime
    )

with c_btn2:
    st.download_button(
        label=f"📄 Скачать Счет ({fmt_code})",
        data=invoice_data,
        file_name=f"Счет_№_{doc_num}_от_{doc_date}.{inv_ext}",
        mime=inv_mime
    )
